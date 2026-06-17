from pathlib import Path

import fitz
from docx import Document as DocxFile
from llama_index.core import Document, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter

from app.config import get_settings
from app.rag.engine import configure_models, get_vector_store


CHUNK_SIZE = 700
CHUNK_OVERLAP = 150

EMERGENCY_MARKERS = ["AIIMS", "incident control", "incident management team"]


def detect_scope(text: str) -> str:
    lowered = text.lower()
    for marker in EMERGENCY_MARKERS:
        if marker.lower() in lowered:
            return "emergency-only"
    return "general"


def load_pdf_pages(documents_dir: Path) -> list[Document]:
    pages = []
    for pdf_path in sorted(documents_dir.glob("*.pdf")):
        pdf = fitz.open(pdf_path)
        for page_index in range(pdf.page_count):
            page_text = pdf[page_index].get_text().strip()
            if not page_text:
                continue
            pages.append(
                Document(
                    text=page_text,
                    metadata={
                        "source": pdf_path.name,
                        "page": page_index + 1,
                        "scope": detect_scope(page_text),
                    },
                )
            )
        pdf.close()
    return pages


def is_heading(paragraph) -> bool:
    style_name = getattr(paragraph.style, "name", "") or ""
    return style_name.startswith("Heading") or style_name == "Title"


def read_docx_sections(docx_path: Path) -> list[tuple[str, str]]:
    docx_file = DocxFile(docx_path)

    sections = []
    current_heading = ""
    current_lines = []

    def flush_current_section():
        joined_text = "\n".join(current_lines).strip()
        if joined_text:
            sections.append((current_heading, joined_text))

    for paragraph in docx_file.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text:
            continue
        if is_heading(paragraph):
            flush_current_section()
            current_heading = paragraph_text
            current_lines = []
        else:
            current_lines.append(paragraph_text)

    flush_current_section()

    table_lines = []
    for table in docx_file.tables:
        for row in table.rows:
            filled_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    filled_cells.append(cell_text)
            if filled_cells:
                table_lines.append(" | ".join(filled_cells))
    if table_lines:
        sections.append(("Tables", "\n".join(table_lines)))

    return sections


def load_docx_documents(documents_dir: Path) -> list[Document]:
    documents = []
    for docx_path in sorted(documents_dir.glob("*.docx")):
        if docx_path.name.startswith("~$"):
            continue
        for heading, section_text in read_docx_sections(docx_path):
            documents.append(
                Document(
                    text=section_text,
                    metadata={
                        "source": docx_path.name,
                        "section": heading,
                        "scope": detect_scope(heading + "\n" + section_text),
                    },
                )
            )
    return documents


def clear_existing_collection():
    settings = get_settings()
    vector_store = get_vector_store()
    client = vector_store.client
    if client.collection_exists(settings.qdrant_collection):
        client.delete_collection(settings.qdrant_collection)


def ingest_documents():
    configure_models()
    settings = get_settings()
    documents_dir = Path(settings.documents_dir)

    clear_existing_collection()

    pdf_pages = load_pdf_pages(documents_dir)
    docx_sections = load_docx_documents(documents_dir)
    all_documents = pdf_pages + docx_sections

    vector_store = get_vector_store()
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    VectorStoreIndex.from_documents(
        all_documents,
        storage_context=storage_context,
        transformations=[splitter],
    )

    print(
        f"Ingested {len(pdf_pages)} PDF pages and {len(docx_sections)} DOCX sections "
        f"into collection '{settings.qdrant_collection}'."
    )


if __name__ == "__main__":
    ingest_documents()
